"""
docx_renderer.py — Dict → Word document via python-docx.

Generates styled Word documents with:
- Document header (title, subtitle, meta)
- Section cards with headings
- Tables with formatting
- Stakeholder cards
- Badges as colored text
- Proper fonts and spacing for CJK content
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Color mapping for badges
BADGE_COLORS = {
    # Stance
    "champion": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "sponsor": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "supporter": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "neutral": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "non-supporter": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "adversary": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "unknown": RGBColor(0x6B, 0x72, 0x80) if HAS_DOCX else None,
    # Role
    "decision-maker": RGBColor(0x4C, 0x1D, 0x95) if HAS_DOCX else None,
    "technical-evaluator": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "economic-buyer": RGBColor(0x92, 0x40, 0x0E) if HAS_DOCX else None,
    "influencer": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "end-user": RGBColor(0x6B, 0x72, 0x80) if HAS_DOCX else None,
    "blocker": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    # Priority
    "high": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "medium": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "low": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "must-meet": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "important": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "nice-to-have": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    # Result / Status
    "achieved": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "partial": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "not-achieved": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "done": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "next": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "planned": RGBColor(0x9C, 0xA3, 0xAF) if HAS_DOCX else None,
    # Action / Gap
    "pending": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "open": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    "in-progress": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "answered": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "unanswered": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
    # Change types (PMR)
    "update": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "add": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "remove": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "no-change": RGBColor(0x6B, 0x72, 0x80) if HAS_DOCX else None,
    "confirm": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "新增": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    "更新": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "删除": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "确认": RGBColor(0x05, 0x96, 0x69) if HAS_DOCX else None,
    # Milestone
    "skipped": RGBColor(0x6B, 0x72, 0x80) if HAS_DOCX else None,
    # Gap
    "still a gap": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    # Objection categories (match DESIGN_SPEC)
    "status quo": RGBColor(0x6B, 0x72, 0x80) if HAS_DOCX else None,
    "status-quo": RGBColor(0x6B, 0x72, 0x80) if HAS_DOCX else None,
    "price/value": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "price/competition": RGBColor(0x6D, 0x28, 0xD9) if HAS_DOCX else None,
    "capability": RGBColor(0xEA, 0x58, 0x0C) if HAS_DOCX else None,
    "risk/trust": RGBColor(0xDC, 0x26, 0x26) if HAS_DOCX else None,
    "authority/process": RGBColor(0xD9, 0x77, 0x06) if HAS_DOCX else None,
}

DOC_TYPE_LABELS = {
    "engagement-plan": "ENGAGEMENT PLAN",
    "call-plan": "CALL PLAN",
    "executive-briefing": "EXECUTIVE BRIEFING",
    "post-meeting-report": "POST-MEETING REPORT",
}


def render_docx(doc: dict, output_path: str) -> str:
    """Render parsed document dict to a Word .docx file.
    
    Args:
        doc: Validated document dict from parse.py
        output_path: Path to write the .docx file
        
    Returns:
        Absolute path to the generated .docx file
    """
    if not HAS_DOCX:
        raise RuntimeError(
            "python-docx is not installed. Install with: pip install python-docx"
        )
    
    output_path = str(Path(output_path).resolve())
    frontmatter = doc.get("frontmatter", {})
    sections = doc.get("sections", [])
    doc_type = frontmatter.get("type", "unknown")
    
    # Create document
    document = Document()
    
    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    
    # Set CJK font
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.append(rFonts)
    
    # EB: Confidential banner
    if doc_type == "executive-briefing":
        _add_confidential_banner(document, frontmatter)
    
    # Document header
    _add_header(document, frontmatter, doc_type)
    
    # Sections
    for i, section in enumerate(sections, 1):
        _add_section(document, section, doc_type, section_num=i)
    
    # Footer
    _add_footer(document, frontmatter, doc_type)
    
    # Save
    document.save(output_path)
    logger.info(f"DOCX generated: {output_path}")
    
    return output_path


def _add_confidential_banner(document, frontmatter):
    """Add confidential banner for EB documents."""
    classification = frontmatter.get("classification", "INTERNAL USE ONLY — AWS Confidential")
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"🔒 {classification}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    _set_paragraph_spacing(para, before=0, after=12)


def _add_header(document, frontmatter, doc_type):
    """Add document header."""
    doc_label = DOC_TYPE_LABELS.get(doc_type, "DOCUMENT")
    customer = frontmatter.get("customer", "")
    
    # Doc type label
    para = document.add_paragraph()
    run = para.add_run(doc_label)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    run.bold = True
    _set_paragraph_spacing(para, before=0, after=2)
    
    # Customer title
    para = document.add_paragraph()
    run = para.add_run(customer)
    run.font.size = Pt(22)
    run.bold = True
    _set_paragraph_spacing(para, before=0, after=4)
    
    # Subtitle
    subtitle_parts = []
    if doc_type in ("call-plan", "executive-briefing", "post-meeting-report"):
        meeting_title = frontmatter.get("meeting_title", "")
        opportunity = frontmatter.get("opportunity", "")
        if meeting_title:
            subtitle_parts.append(meeting_title)
        if opportunity:
            subtitle_parts.append(opportunity)
    elif doc_type == "engagement-plan":
        opportunity = frontmatter.get("opportunity", "")
        tcv = frontmatter.get("tcv", "")
        if opportunity:
            subtitle_parts.append(opportunity)
        if tcv:
            subtitle_parts.append(f"{tcv} TCV")
    
    if subtitle_parts:
        para = document.add_paragraph()
        run = para.add_run(" · ".join(subtitle_parts))
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _set_paragraph_spacing(para, before=0, after=6)
    
    # Meta line
    meta_parts = []
    stage = frontmatter.get("stage", "")
    if stage:
        meta_parts.append(f"📈 {stage}")
    date = frontmatter.get("date", frontmatter.get("created", ""))
    if date:
        meta_parts.append(f"📅 {date}")
    
    if meta_parts:
        para = document.add_paragraph()
        run = para.add_run("  |  ".join(meta_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _set_paragraph_spacing(para, before=0, after=12)
    
    # Separator
    document.add_paragraph("─" * 60)


def _add_section(document, section, doc_type, section_num=None):
    """Add a section to the document."""
    emoji = section.get("emoji", "")
    title = section.get("title", "")
    content_blocks = section.get("content", [])
    
    # Section heading: "N. emoji Title" format (consistent with HTML/PDF)
    emoji_str = f"{emoji} " if emoji else ""
    num_prefix = f"{section_num}. " if section_num else ""
    heading_text = f"{num_prefix}{emoji_str}{title}"
    heading = document.add_heading(heading_text, level=2)
    if heading.runs:
        heading.runs[0].font.size = Pt(14)
    
    # Content blocks
    for block in content_blocks:
        _add_block(document, block, doc_type)


def _add_block(document, block, doc_type):
    """Add a content block to the document."""
    block_type = block.get("type", "")
    
    if block_type == "table":
        _add_table(document, block)
    elif block_type == "stakeholder_card":
        _add_stakeholder_card(document, block, doc_type)
    elif block_type == "milestone":
        _add_milestone(document, block)
    elif block_type == "objection":
        _add_objection(document, block)
    elif block_type == "eb_person":
        _add_eb_person(document, block)
    elif block_type == "subsection":
        _add_subsection(document, block, doc_type)
    elif block_type == "subsubsection":
        _add_subsubsection(document, block, doc_type)
    elif block_type == "bullet_list":
        _add_bullet_list(document, block)
    elif block_type == "paragraph":
        _add_paragraph(document, block)
    elif block_type == "progress_bar":
        _add_progress_bar(document, block)
    elif block_type == "highlight":
        _add_highlight(document, block)
    elif block_type == "objective":
        _add_objective(document, block)
    elif block_type == "concern":
        _add_concern(document, block)
    elif block_type == "engagement_entry":
        _add_engagement_entry(document, block)


def _add_table(document, block):
    """Add a markdown table to the document."""
    headers = block.get("headers", [])
    rows = block.get("rows", [])
    
    if not headers and not rows:
        return
    
    # Determine actual column count (rows may be wider than headers)
    col_count = max(len(headers), *(len(r) for r in rows)) if rows else len(headers)
    
    # Pad headers if rows are wider (e.g. empty first header stripped by parser)
    if headers and len(headers) < col_count:
        headers = [""] * (col_count - len(headers)) + headers
    elif not headers:
        headers = [""] * col_count
    
    table = document.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _style_cell(cell, bold=True, size=9, color=RGBColor(0x6B, 0x72, 0x80))
    
    # Data rows
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            if col_idx < col_count:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = _strip_markdown(cell_text)
                _style_cell(cell, size=10)
    
    document.add_paragraph()  # spacing


def _add_stakeholder_card(document, block, doc_type):
    """Add a stakeholder card."""
    name = block.get("name", "")
    title = block.get("title", "")
    stance = block.get("stance", "")
    role = block.get("role", "")
    
    # Name + badges
    para = document.add_paragraph()
    run = para.add_run(f"{name}")
    run.bold = True
    run.font.size = Pt(12)
    
    if stance:
        run = para.add_run(f"  [{stance.title()}]")
        color = BADGE_COLORS.get(stance.lower())
        if color:
            run.font.color.rgb = color
        run.font.size = Pt(9)
    
    if role:
        run = para.add_run(f"  [{role.replace('-', ' ').title()}]")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    # Title
    if title:
        para = document.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _set_paragraph_spacing(para, before=0, after=4)
    
    # Fields
    if doc_type == "call-plan":
        _add_field(document, "Focus", block.get("focus", ""))
        _add_field(document, "Communication", block.get("communication", ""))
        goal = block.get("our_goal", "") or block.get("how_to_win", "")
        if goal:
            _add_field(document, "🎯 Our Goal", goal, highlight=True)
    else:
        _add_field(document, "What They Care About", block.get("what_they_care_about", ""))
        _add_field(document, "What We Need", block.get("what_we_need", ""))
        win = block.get("how_to_win", "") or block.get("our_goal", "")
        if win:
            _add_field(document, "🎯 How to Win", win, highlight=True)
    
    document.add_paragraph()  # spacing


def _add_milestone(document, block):
    """Add a milestone entry."""
    number = block.get("number", 0)
    title = block.get("title", "")
    status = block.get("status", "planned")
    timeline = block.get("timeline", "")
    
    status_icons = {"done": "✓", "next": "▶", "planned": "○"}
    icon = status_icons.get(status, "○")
    
    para = document.add_paragraph()
    run = para.add_run(f"{icon} Milestone {number}: {title}")
    run.bold = True
    run.font.size = Pt(11)
    
    color = BADGE_COLORS.get(status)
    if color:
        run.font.color.rgb = color
    
    # Meta
    meta_parts = []
    if block.get("stakeholders"):
        meta_parts.append(f"👤 {block['stakeholders']}")
    if block.get("aws_resources"):
        meta_parts.append(f"🏢 {block['aws_resources']}")
    if timeline:
        meta_parts.append(f"📅 {timeline}")
    
    if meta_parts:
        para = document.add_paragraph()
        run = para.add_run("  |  ".join(meta_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        _set_paragraph_spacing(para, before=0, after=8)


def _add_objection(document, block):
    """Add an objection card."""
    title = block.get("title", "")
    category = block.get("category", "")
    likely_from = block.get("likely_from", "")
    response = block.get("response", "")
    plan_b = block.get("plan_b", "")
    
    para = document.add_paragraph()
    run = para.add_run(f"⚡ {_strip_markdown(title)}")
    run.bold = True
    run.font.size = Pt(11)
    
    if category:
        run = para.add_run(f"  [{category.replace('-', ' ').title()}]")
        run.font.size = Pt(9)
        color = BADGE_COLORS.get(category.split('-')[0] if '-' in category else category)
        if color:
            run.font.color.rgb = color
    
    if likely_from:
        para = document.add_paragraph()
        run = para.add_run(f"Likely From: {likely_from}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    if response:
        _add_field(document, "Response", response)
    if plan_b:
        _add_field(document, "Plan B", plan_b)
    
    document.add_paragraph()


def _add_eb_person(document, block):
    """Add an EB person background."""
    name = block.get("name", "")
    title = block.get("title", "")
    stance = block.get("stance", "")
    fields = block.get("fields", {})
    
    para = document.add_paragraph()
    run = para.add_run(f"{name} — {title}")
    run.bold = True
    run.font.size = Pt(12)
    
    if stance and stance != "unknown":
        run = para.add_run(f"  [{stance.title()}]")
        color = BADGE_COLORS.get(stance.lower())
        if color:
            run.font.color.rgb = color
        run.font.size = Pt(9)
    
    for label, value in fields.items():
        _add_field(document, label, value)
    
    document.add_paragraph()


def _add_subsection(document, block, doc_type):
    """Add a subsection."""
    title = block.get("title", "")
    emoji = block.get("emoji", "")
    content = block.get("content", [])
    
    prefix = f"{emoji} " if emoji else ""
    heading = document.add_heading(f"{prefix}{title}", level=3)
    if heading.runs:
        heading.runs[0].font.size = Pt(12)
    
    for b in content:
        _add_block(document, b, doc_type)


def _add_subsubsection(document, block, doc_type="unknown"):
    """Add a sub-subsection (#### level)."""
    title = block.get("title", "")
    emoji = block.get("emoji", "")
    content = block.get("content", [])

    prefix = f"{emoji} " if emoji else ""
    heading = document.add_heading(f"{prefix}{title}", level=4)
    if heading.runs:
        heading.runs[0].font.size = Pt(11)

    for b in content:
        _add_block(document, b, doc_type)


def _add_bullet_list(document, block):
    """Add a bullet list."""
    for item in block.get("items", []):
        para = document.add_paragraph(style='List Bullet')
        run = para.add_run(_strip_markdown(item))
        run.font.size = Pt(10)


def _add_paragraph(document, block):
    """Add a paragraph."""
    text = block.get("text", "")
    if text:
        para = document.add_paragraph()
        run = para.add_run(_strip_markdown(text))
        run.font.size = Pt(10)


def _add_progress_bar(document, block):
    """Add a metro-style progress bar as a table for proper alignment."""
    from docx.table import _Cell
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    stages = block.get("stages", [])
    current = block.get("current", 0)
    if not stages:
        return
    
    n = len(stages)
    # Build a 2-row table: row 0 = node symbols + connectors, row 1 = labels
    # Columns: stage, connector, stage, connector, ..., stage
    col_count = n * 2 - 1  # stages + connectors interleaved
    
    table = document.add_table(rows=2, cols=col_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Remove table borders
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    # Set no borders explicitly
    from docx.oxml import OxmlElement
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        borders.append(el)
    tblPr.append(borders)
    
    for idx in range(n):
        col_idx = idx * 2  # stage columns at 0, 2, 4, ...
        
        # Row 0: Node symbol
        cell = table.cell(0, col_idx)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx < current:
            symbol, color = "●", RGBColor(0x05, 0x96, 0x69)
            size = Pt(13)
        elif idx == current:
            symbol, color = "◉", RGBColor(0x6D, 0x28, 0xD9)
            size = Pt(16)
        else:
            symbol, color = "○", RGBColor(0x9C, 0xA3, 0xAF)
            size = Pt(13)
        
        run = cell.paragraphs[0].add_run(symbol)
        run.font.size = size
        run.font.color.rgb = color
        run.bold = (idx == current)
        
        # Row 1: Label
        label_cell = table.cell(1, col_idx)
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx < current:
            lbl_color = RGBColor(0x05, 0x96, 0x69)
        elif idx == current:
            lbl_color = RGBColor(0x6D, 0x28, 0xD9)
        else:
            lbl_color = RGBColor(0x9C, 0xA3, 0xAF)
        
        lbl_run = label_cell.paragraphs[0].add_run(stages[idx])
        lbl_run.font.size = Pt(7)
        lbl_run.font.color.rgb = lbl_color
        lbl_run.bold = (idx == current)
        
        # Connector column (between stages)
        if idx < n - 1:
            conn_col = col_idx + 1
            conn_cell = table.cell(0, conn_col)
            conn_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            conn_color = RGBColor(0x05, 0x96, 0x69) if idx < current else RGBColor(0xD1, 0xD5, 0xDB)
            conn_run = conn_cell.paragraphs[0].add_run("━━━")
            conn_run.font.size = Pt(9)
            conn_run.font.color.rgb = conn_color


def _add_highlight(document, block):
    """Add a highlight/callout box."""
    text = block.get("text", "")
    if text:
        para = document.add_paragraph()
        run = para.add_run(_strip_markdown(text))
        run.font.size = Pt(10)
        run.bold = True


def _add_objective(document, block):
    """Add an EB objective."""
    title = block.get("title", "")
    context = block.get("context", "")
    talking_points = block.get("talking_points", [])
    ask = block.get("ask", "")
    
    para = document.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    
    if context:
        _add_field(document, "Context", context)
    
    if talking_points and isinstance(talking_points, list):
        para = document.add_paragraph()
        run = para.add_run("Talking Points:")
        run.bold = True
        run.font.size = Pt(10)
        for tp in talking_points:
            p = document.add_paragraph(style='List Bullet')
            r = p.add_run(_strip_markdown(tp))
            r.font.size = Pt(10)
    
    if ask:
        _add_field(document, "Ask", ask, highlight=True)
    
    document.add_paragraph()


def _add_concern(document, block):
    """Add an EB concern."""
    title = block.get("title", "")
    
    para = document.add_paragraph()
    run = para.add_run(f"⚠️ {_strip_markdown(title)}")
    run.bold = True
    run.font.size = Pt(11)
    
    _add_field(document, "Acknowledge", block.get("acknowledge", ""))
    _add_field(document, "Pivot", block.get("pivot", ""))
    _add_field(document, "Elevate", block.get("elevate", ""))
    if block.get("landmine"):
        _add_field(document, "💣 Landmine", block.get("landmine", ""))
    
    document.add_paragraph()


def _add_engagement_entry(document, block):
    """Add an engagement log entry."""
    number = block.get("number", 0)
    date = block.get("date", "")
    
    para = document.add_paragraph()
    run = para.add_run(f"Engagement #{number} — {date}")
    run.bold = True
    run.font.size = Pt(11)
    
    fields = [
        ("Attendees", block.get("attendees", "")),
        ("Planned", block.get("planned", "")),
        ("Actual", block.get("actual", "")),
        ("People Updates", block.get("people_updates", "")),
        ("Key Learnings", block.get("key_learnings", "")),
        ("Plan Adjustment", block.get("plan_adjustment", "")),
    ]
    
    for label, value in fields:
        if value:
            _add_field(document, label, value)
    
    document.add_paragraph()


def _add_field(document, label, value, highlight=False):
    """Add a labeled field with provenance label support."""
    if not value:
        return
    para = document.add_paragraph()
    run = para.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(10)
    _add_text_with_provenance(para, _strip_markdown(value), size=10)
    if highlight:
        # Color the last run
        for run in para.runs:
            if run.text and run.text != f"{label}: ":
                run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    _set_paragraph_spacing(para, before=0, after=2)


def _add_footer(document, frontmatter, doc_type):
    """Add document footer."""
    document.add_paragraph("─" * 60)
    
    doc_label = DOC_TYPE_LABELS.get(doc_type, "DOCUMENT")
    customer = str(frontmatter.get("customer", ""))
    date = str(frontmatter.get("date", frontmatter.get("version", "")))
    
    footer_parts = [doc_label, customer]
    if date:
        footer_parts.append(date)
    
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(" | ".join(footer_parts))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Generated by Dali · Confidential")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)


def _set_paragraph_spacing(para, before=0, after=6):
    """Set paragraph spacing in points."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _style_cell(cell, bold=False, size=10, color=None):
    """Style a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color


def _strip_markdown(text: str) -> str:
    """Strip markdown formatting from text for Word (preserves provenance labels)."""
    import re
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove badge syntax
    text = re.sub(r'\{(\w+):([\w-]+)\}', r'\2', text)
    return text.strip()


# Provenance label colors (RGB)
_PROVENANCE_COLORS = {
    '[销售确认]': RGBColor(0x16, 0x65, 0x34),  # green-800
    '[网络搜索]': RGBColor(0x1E, 0x40, 0xAF),  # blue-800
    '[AI推断]': RGBColor(0x6B, 0x72, 0x80),    # gray-500
}


def _add_text_with_provenance(para, text: str, size=10, bold=False):
    """Add text to a paragraph, rendering provenance labels as colored runs."""
    import re
    # Split on provenance labels, keeping the delimiters
    parts = re.split(r'(\[销售确认\]|\[网络搜索\]|\[AI推断\])', text)
    for part in parts:
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = Pt(size)
        run.bold = bold
        if part in _PROVENANCE_COLORS:
            run.font.color.rgb = _PROVENANCE_COLORS[part]
            run.font.size = Pt(max(7, size - 2))
            run.bold = False
